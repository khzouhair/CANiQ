import os
import shutil
from typing import List, Dict, Optional
from dotenv import load_dotenv
from tqdm import tqdm
from docx import Document as DocxDocument
import chardet
import re
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
import google.generativeai as genai
from typing import TYPE_CHECKING
import json
import concurrent.futures
import fitz 
from functools import lru_cache
import logging


load_dotenv()

if TYPE_CHECKING:
    import tiktoken 
try:
    import tiktoken
except Exception:
    tiktoken = None

############################
# Configs (env or defaults)
############################
DATA_DIR = os.getenv("DATA_DIR", "../data/raw_documents")
FAISS_DIR = os.getenv("FAISS_DIR", "../data/vector_store")
EMBED_MODEL = os.getenv("EMBED_MODEL", "all-MiniLM-L6-v2.gguf2.f16.gguf")
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY", "AIzaSyDNIMaOahHZyp2TjX1qCbPCQmdlZWv_3CM")
CHUNK_SIZE = int(os.getenv("CHUNK_SIZE", 500))
CHUNK_OVERLAP = int(os.getenv("CHUNK_OVERLAP", 120))
TOP_K = int(os.getenv("TOP_K", 4))

gpt4all_embeddings = None

modele_name = "gemini-2.5-flash"
#########################
# 1) Parsing / Extraction
#########################
def _is_temp_office_file(path: str) -> bool:
    return os.path.basename(path).startswith("~$")

def _extract_text_worker(path: str) -> Optional[Dict]:
    if _is_temp_office_file(path):
        logging.info(f"[extract_text_worker] fichier temporaire Office ignoré: {path}")
        return None
    ext = os.path.splitext(path)[1].lower()
    try:
        if ext == ".pdf":
            doc = fitz.open(path)
            pages = [p.get_text("text") or "" for p in doc]
            text = "\n".join(pages)
            doc.close()
        elif ext == ".docx":
            try:
                doc = DocxDocument(path)
                text = "\n".join([p.text for p in doc.paragraphs])
            except Exception as e:
                logging.exception(f"[extract_text_worker] docx unreadable {path}: {e}")
                return None
        else:
            with open(path, "rb") as f:
                raw = f.read()
            enc = chardet.detect(raw).get("encoding") or "utf-8"
            text = raw.decode(enc, errors="ignore")
        return {"path": path, "text": text, "size": os.path.getsize(path), "mtime": os.path.getmtime(path)}
    except Exception as e:
        logging.exception(f"[extract_text_worker] erreur pour {path}: {e}")
        return None

def load_all_files_parallel(
    base_path: str = DATA_DIR,
    allowed_ext=(".pdf", ".docx", ".txt"),
    include_empty: bool = False,
    max_workers: int = None,
    use_cache: bool = False
) -> List[Dict]:
    base_path = os.path.abspath(base_path)
    manifest_path = os.path.join(base_path, ".extract_manifest.json")
    prev_manifest = {}
    if use_cache and os.path.exists(manifest_path):
        try:
            with open(manifest_path, "r", encoding="utf-8") as f:
                prev_manifest = json.load(f)
        except Exception:
            prev_manifest = {}

    paths = []
    for root, _, files in os.walk(base_path):
        for fn in files:
            ext = os.path.splitext(fn)[1].lower()
            if ext in allowed_ext:
                paths.append(os.path.join(root, fn))
    if not paths:
        logging.info("Aucun fichier trouvé.")
        return []

    max_workers = max_workers or min( (os.cpu_count() or 4), 8)
    results = []
    # Extraction en processus séparés (meilleur pour le parsing PDF lié au CPU)
    with concurrent.futures.ProcessPoolExecutor(max_workers=max_workers) as ex:
        for res in tqdm(ex.map(_extract_text_worker, paths), total=len(paths)):
            if not res:
                continue
            rel = os.path.relpath(res["path"], base_path).replace(os.sep, "/")
            key = rel
            # cache check
            if use_cache:
                prev = prev_manifest.get(key)
                if prev and prev.get("mtime") == res["mtime"] and prev.get("size") == res["size"]:
                    # no change -> can skip or reload text from prev storage
                    logging.info(f"[load_all_files_parallel] inchangé, ignoré: {rel}")
                    continue
            if not include_empty and (not res["text"] or res["text"].strip() == ""):
                logging.info(f"[load_all_files_parallel] fichier vide ignoré: {rel}")
                continue
            results.append({
                "source": rel,
                "text": res["text"],
                "size": res["size"],
                "mtime": res["mtime"]
            })

    
    if use_cache:
        new_manifest = { r["source"]: {"mtime": r["mtime"], "size": r["size"]} for r in results }
        try:
            with open(manifest_path, "w", encoding="utf-8") as f:
                json.dump(new_manifest, f, indent=2)
        except Exception:
            logging.exception("Impossible d'écrire le manifest de cache.")

    logging.info(f"[load_all_files_parallel] {len(results)} fichiers extraits depuis {base_path}")
    return results

##############################
# 2) Cleaning & Normalization
##############################
import unicodedata

def clean_text(text: str) -> str:

    if not text or not isinstance(text, str):
        return ""

    # Normalisation unicode
    t = unicodedata.normalize("NFKC", text)

    # unifier les séparateurs de nouvelle ligne
    t = t.replace("\r\n", "\n").replace("\r", "\n")
    t = t.replace("\u2028", "\n").replace("\u2029", "\n")

    # normaliser les tirets et supprimer le trait d'union souple
    t = t.replace("\u00ad", "")

    # Réduire les lignes vides excessives
    t = re.sub(r"\n\s*\n+", "\n\n", t)

    # convertir ":" ou ";" suivi d'un saut de ligne en texte courant (garder un espace après)
    t = re.sub(r"([:;])\s*\n\s*", r"\1 ", t)

    # corriger la césure à travers les lignes
    t = re.sub(r"([A-Za-zÀ-ÖØ-öø-ÿ])-\n\s*([A-Za-zÀ-ÖØ-öø-ÿ])", r"\1 \2", t)

    # supprimer les caractères de contrôle sauf le saut de ligne
    t = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]", "", t)

    # garder uniquement les caractères raisonnables 
    t = re.sub(r"[^A-Za-zÀ-ÖØ-öø-ÿ0-9\s\.,;:'\"\?\!\-\(\)/\n€£$¥]", " ", t)

    # Séparer les lettres et les chiffres s'ils sont concaténés
    t = re.sub(r"([A-Za-zÀ-ÖØ-öø-ÿ]{2,})(?=\d)", r"\1 ", t)
    t = re.sub(r"(\d)(?=[A-Za-zÀ-ÖØ-öø-ÿ]{2,})", r"\1 ", t)

    # remplacer les sauts de ligne simples (pas les sauts de paragraphe) par un espace, avec des heuristiques
    def _nl_repl(m):
        i = m.start()
        prev = t[i - 1] if i > 0 else ""
        following = t[i + 1 : i + 80]
        # garder le saut de ligne si ponctuation de fin de phrase réelle
        if prev and prev in ".!?)]…":
            return "\n"
        # garder le saut de ligne si la ligne suivante ressemble à une liste ou un élément numéroté
        if re.match(r"\s*([-•*]|\d+[.)])", following):
            return "\n"
        return " "

    t = re.sub(r"(?<!\n)\n(?!\n)", _nl_repl, t)

    # réduire la ponctuation répétée comme "-----" ou "???" -> un seul
    t = re.sub(r"([^\w\s])\1{2,}", r"\1", t)

    # normaliser les espaces/onglets
    t = re.sub(r"[ \t]{2,}", " ", t)

    # normaliser les guillemets typographiques
    t = (t.replace("’", "'")
           .replace("“", '"')
           .replace("”", '"')
           .replace("«", '"')
           .replace("»", '"'))

    # s'assurer que les éléments de liste commencent sur une nouvelle ligne 
    t = re.sub(r"(?<!\n)(\s*)([-•*]|\d+[.)])\s", r"\n\1\2 ", t)

    # supprimer les espaces en début et fin de ligne mais préserver les doubles sauts de ligne (paragraphes)
    lines = [ln.strip() for ln in t.split("\n")]
    t = "\n".join(lines)
    t = re.sub(r"\n{3,}", "\n\n", t).strip()

    # Séparer les tokens CamelCase ou alphanumériques mixtes si nécessaire.
    t = re.sub(r"([a-zà-öø-ÿ])(?=[A-ZÀ-Ö]{2,})", r"\1 ", t)       
    t = re.sub(r"([A-Za-zÀ-ÖØ-öø-ÿ]{2,})(?=\d)", r"\1 ", t)       
    t = re.sub(r"([a-zà-öø-ÿ])(?=[A-ZÀ-Ö][a-zà-öø-ÿ])", r"\1 ", t)          

    # Corriger l'espacement de la ponctuation
    t = re.sub(r"\s+([.,])", r"\1", t)  
    t = re.sub(r"([^\s])([!?:;])", r"\1 \2", t)   

    # Réduire les sauts de ligne accidentels avant les listes
    t = re.sub(r"\n{2,}(\s*(?:[-•*]|\d+[.)]))", r"\n\1", t)

    # nettoyage final 
    t = re.sub(r"[ \t]{2,}", " ", t) 
    t = re.sub(r" *\n{2,} *", "\n\n", t).strip() 

    return t

###############
# 3) Chunking
###############
def create_langchain_documents(prepped_docs: List[Dict]) -> List[Document]:
    return [Document(page_content=d["text"], metadata={"source": d["source"]}) for d in prepped_docs]

@lru_cache(maxsize=8192)
def _approx_token_len_cached(s: str) -> int:
    if not s:
        return 0
    # utiliser le tokenizer pour les chaînes de taille raisonnable ; revenir à une approximation rapide pour les gros blocs
    if tiktoken is not None and len(s) <= 5000:
        try:
            try:
                enc = tiktoken.encoding_for_model("gpt-4o-mini")
            except Exception:
                enc = tiktoken.get_encoding("cl100k_base")
            return len(enc.encode(s))
        except Exception:
            pass
    if len(s) < 200:
        return max(1, len(s.split()))
    # heuristique : ~4 caractères par token (empirique) pour les longs textes
    return max(1, int(len(s) / 4))

def _approx_token_len(s: str) -> int:
    # éviter de mettre en cache de très grandes chaînes comme clés
    if not s:
        return 0
    if len(s) > 20000:
        return int(len(s) / 4)
    return _approx_token_len_cached(s)

def chunk_documents(lc_documents: List[Document], chunk_size: int = CHUNK_SIZE, chunk_overlap: int = CHUNK_OVERLAP) -> List[Document]:
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        length_function=_approx_token_len,
        add_start_index=True
    )
    chunks = splitter.split_documents(lc_documents)
    for i, c in enumerate(chunks):
        md = dict(c.metadata) if c.metadata else {}
        md["chunk_index"] = i
        c.metadata = md
    print(f"[chunk_documents] {len(chunks)} chunks générés.")
    return chunks

##############################
# 4) Embeddings & Vector Store 
##############################

def get_hf_embeddings(model_name: str = EMBED_MODEL):
    mdl = model_name or ""
    if ".gguf" in mdl or mdl.endswith(".gguf") or ".gguf2" in mdl:
        logging.warning("EMBED_MODEL semble être un fichier GGUF (%s). Sentence-Transformers ne peut pas le charger. "
                        "Fallback vers 'sentence-transformers/all-MiniLM-L6-v2'. Pour utiliser gguf, utilisez GPT4All/GGML pipeline.", mdl)
        mdl = "sentence-transformers/all-MiniLM-L6-v2"
    try:
        from langchain_community.embeddings import HuggingFaceEmbeddings
    except Exception:
        try:
            from langchain.embeddings import HuggingFaceEmbeddings 
        except Exception:
            HuggingFaceEmbeddings = None

    if HuggingFaceEmbeddings is not None:
        try:
            return HuggingFaceEmbeddings(model_name=mdl)
        except Exception as e:
            logging.warning("La construction de HuggingFaceEmbeddings a échoué : %s. Retour à l'utilisation directe de sentence-transformers.", e)

    try:
        from sentence_transformers import SentenceTransformer
    except Exception as e:
        raise RuntimeError("sentence-transformers non installé ou indisponible. Installez via: pip install sentence-transformers") from e

    try:
        st = SentenceTransformer(mdl)
        class _STWrapper:
            def __init__(self, model):
                self._m = model
            def embed_documents(self, texts):
                return [list(v) for v in self._m.encode(texts, show_progress_bar=True)]
            def embed_query(self, text):
                return list(self._m.encode([text])[0])
        return _STWrapper(st)
    except Exception as e:
        raise RuntimeError(f"Impossible de charger le modèle sentence-transformers '{mdl}': {e}\n"
                           "Vérifiez que EMBED_MODEL est un identifiant HuggingFace valide (ex: 'sentence-transformers/all-MiniLM-L6-v2') "
                           "ou installez/autorisez l'accès au repo privé (hf auth login).") from e

# FAISS vectorstore
try:
        from langchain_community.vectorstores import FAISS
except Exception as e:
        raise RuntimeError("FAISS vectorstore non disponible (pip install faiss-cpu langchain_community) : " + str(e))


def build_faiss_from_chunks(
    chunks: List[Document],
    faiss_dir: str = FAISS_DIR,
    persist: bool = True,
    reset: bool = True,
) -> "FAISS":
    emb = get_hf_embeddings()

    if reset and os.path.exists(faiss_dir):
        shutil.rmtree(faiss_dir)

    
    try:
        db = FAISS.from_documents(chunks, embedding=emb)
    except TypeError:
        db = FAISS.from_documents(chunks, embedding_function=emb)

    if persist:
        os.makedirs(faiss_dir, exist_ok=True)
        db.save_local(faiss_dir)

    return db

def load_faiss(faiss_dir: str = FAISS_DIR, allow_dangerous_deserialization: bool = False):
    try:
        try:
            from langchain_community.vectorstores import FAISS
        except Exception:
            from langchain_community.vectorstores import FAISS
    except Exception as e:
        raise RuntimeError("FAISS non disponible : " + str(e))

    if not os.path.exists(faiss_dir):
        return None

    emb = get_hf_embeddings()
    return FAISS.load_local(faiss_dir, embeddings=emb, allow_dangerous_deserialization=allow_dangerous_deserialization)

##########################
# 5) Retriever & RAG Chain 
##########################

def build_simple_retriever(db):
    retriever = db.as_retriever(search_type="similarity", search_kwargs={"k": TOP_K})
    return retriever


def generate_answer_from_chunks(question: str, chunks: list[Document]):
    context_text = "\n\n---\n\n".join(
        [
            f"Source: {c.metadata.get('source', 'unknown')} | chunk_index: {i}\n{c.page_content.strip()}"
            for i, c in enumerate(chunks)
        ]
    )
    prompt = (
    '''Tu es un assistant expert sur la CAN 2025, précis, concis et fiable en français. 
Utilise UNIQUEMENT le contexte ci‑dessous pour répondre. Ne fais AUCUNE supposition. 
- Donne des réponses complètes, structurées et claires. 
- Inclue tous les détails pertinents : équipes, joueurs, stades, scores, horaires, règlements, statistiques officielles. 
- Limite la réponse à 6‑10 phrases. 
- Liste toutes les informations pertinentes sans répétition inutile.\n\n'''
    f"Contexte:\n{context_text}\n\nQuestion: {question}\n\nRéponse (6 à 10 phrases) :"
)

    genai.configure(api_key=GEMINI_API_KEY)

    try:
        model = genai.GenerativeModel(modele_name)

        response = model.generate_content(prompt)

        return response.text.strip(), prompt

    except Exception as e:
        return f"[Erreur Gemini] {e}", prompt


######################### 
# 6) High-level pipelines 
#########################
import logging

def full_preprocess_and_index(
    data_dir: str = DATA_DIR,
    faiss_dir: str = FAISS_DIR,
    reset: bool = True,
    persist: bool = True,
    max_workers: Optional[int] = None,
    use_cache: bool = False,
) -> int:
    logging.info("full_preprocess_and_index: scanning %s", data_dir)

    raw_docs = load_all_files_parallel(data_dir, max_workers=max_workers, use_cache=use_cache)
    prepped = []
    for d in raw_docs:
        txt = d.get("text", "") or ""
        cleaned = clean_text(txt)
        if cleaned:
            prepped.append({"source": d["source"], "text": cleaned})

    lc_docs = create_langchain_documents(prepped)
    chunks = chunk_documents(lc_docs)

    logging.info("full_preprocess_and_index: building FAISS in %s (%d chunks)", faiss_dir, len(chunks))
    build_faiss_from_chunks(chunks, faiss_dir=faiss_dir, persist=persist, reset=reset)
    return len(chunks)


def _invoke_retriever(retriever, question: str, db, top_k: int):
    
    try:
        if hasattr(retriever, "get_relevant_documents"):
            return retriever.get_relevant_documents(question)
    except Exception:
        pass

    try:
        out = retriever(question)
        if isinstance(out, list):
            return out
        if isinstance(out, dict):
            for k in ("output", "results", "documents"):
                if k in out and isinstance(out[k], list):
                    return out[k]
    except Exception:
        pass

    try:
        if hasattr(retriever, "run"):
            out = retriever.run(question)
            if isinstance(out, list):
                return out
    except Exception:
        pass

    try:
        return db.as_retriever(search_type="similarity", search_kwargs={"k": top_k}).invoke(question)
    except Exception as e:
        raise RuntimeError("Impossible d'invoquer le retriever: " + str(e))


def retrieve_top_k(
    question: str,
    faiss_dir: str = FAISS_DIR,
    top_k: int = TOP_K,
    allow_dangerous_deserialization: bool = False,
):
    db = load_faiss(faiss_dir, allow_dangerous_deserialization=allow_dangerous_deserialization)
    if db is None:
        raise FileNotFoundError(f"Vector store introuvable dans {faiss_dir}. Lance full_preprocess_and_index() d'abord.")
    retriever = build_simple_retriever(db)
    docs = _invoke_retriever(retriever, question, db, top_k)
    return docs[:top_k]


def answer_question_flow(
    question: str,
    faiss_dir: str = FAISS_DIR,
    top_k: int = TOP_K,
    allow_dangerous_deserialization: bool = False,
    chat_history: Optional[List[Dict[str, str]]] = None) -> dict:
    
    docs = retrieve_top_k(question, faiss_dir=faiss_dir, top_k=top_k, allow_dangerous_deserialization=allow_dangerous_deserialization)
    if not docs:
        return {"answer": "Aucune information trouvée dans le corpus.", "sources": [], "prompt": ""}
    answer, prompt_used = generate_answer_from_chunks(question, docs)
    sources = []
    for d in docs:
        src = None
        text = ""
        if hasattr(d, "metadata"):
            src = d.metadata.get("source")
        elif isinstance(d, dict):
            src = d.get("metadata", {}).get("source")
        if hasattr(d, "page_content"):
            text = d.page_content or ""
        elif isinstance(d, dict):
            text = d.get("page_content", "") or d.get("text", "")
        sources.append({"source": src or "unknown", "text": (text or "")[:2000]})
    return {"answer": answer, "sources": sources, "prompt": prompt_used}

